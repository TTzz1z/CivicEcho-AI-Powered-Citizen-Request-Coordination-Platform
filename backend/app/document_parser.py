"""Document parsing for KB uploads.

Supports PDF, Word (.docx), Markdown and plain text. Detects scanned PDFs and
marks OCR status/quality. PDFs that look scanned are flagged with
``ocr_status='required'`` but text extraction still attempts OCR-free path;
this keeps the dependency footprint small (no Tesseract runtime required) while
still recording the OCR signal for admins.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    file_type: str
    ocr_status: str  # none / required / done / failed
    ocr_quality: float | None  # 0.0 - 1.0
    page_count: int | None
    parser_notes: list[str]

    @property
    def ok(self) -> bool:
        return bool(self.text and self.text.strip())


def parse_bytes(data: bytes, filename: str, mime_type: str | None = None) -> ParsedDocument:
    """Parse uploaded bytes into text. Returns ParsedDocument."""
    name = (filename or "").lower()
    if name.endswith(".pdf") or (mime_type or "").lower() == "application/pdf":
        return _parse_pdf(data)
    if name.endswith(".docx") or (mime_type or "").lower() in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        return _parse_docx(data)
    if name.endswith(".doc") or (mime_type or "").lower() == "application/msword":
        # python-docx does not support legacy .doc; flag as failed.
        return ParsedDocument(text="", file_type="word", ocr_status="none",
                              ocr_quality=None, page_count=None,
                              parser_notes=["legacy .doc 不支持，请转换为 .docx 后再上传"])
    if name.endswith((".md", ".markdown")):
        return _parse_markdown(data.decode("utf-8", errors="ignore"))
    if name.endswith((".txt", ".text")) or (mime_type or "").lower().startswith("text/"):
        return ParsedDocument(
            text=data.decode("utf-8", errors="ignore"),
            file_type="text", ocr_status="none", ocr_quality=None,
            page_count=None, parser_notes=[],
        )
    # Fallback: try utf-8 decode
    return ParsedDocument(
        text=data.decode("utf-8", errors="ignore"),
        file_type="text", ocr_status="none", ocr_quality=None,
        page_count=None, parser_notes=["未识别的文件类型，按纯文本处理"],
    )


def _parse_pdf(data: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        logger.warning("pypdf not available: %s", exc)
        return ParsedDocument(text="", file_type="pdf", ocr_status="failed",
                              ocr_quality=None, page_count=None,
                              parser_notes=["pypdf 未安装"])
    try:
        reader = PdfReader(io.BytesIO(data))
        page_count = len(reader.pages)
        pages_text: list[str] = []
        non_text_pages = 0
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception as exc:  # extraction can fail per-page
                logger.debug("page extract failed: %s", exc)
                txt = ""
            pages_text.append(txt)
            # Heuristic: a page with very few extractable chars likely is scanned
            stripped = re.sub(r"\s+", "", txt)
            if len(stripped) < 20:
                non_text_pages += 1
        text = "\n\n".join(pages_text).strip()
        # OCR signal: if >50% of pages have no extractable text, mark as scanned
        if page_count > 0 and non_text_pages / page_count > 0.5:
            ocr_status = "required"
            ocr_quality = round(1.0 - non_text_pages / page_count, 3)
            notes = [f"检测到 {non_text_pages}/{page_count} 页疑似扫描内容，建议补充 OCR 文本或上传文本版本"]
        else:
            ocr_status = "none"
            ocr_quality = None
            notes = []
        return ParsedDocument(
            text=text, file_type="pdf", ocr_status=ocr_status,
            ocr_quality=ocr_quality, page_count=page_count, parser_notes=notes,
        )
    except Exception as exc:
        logger.warning("PDF parse failed: %s", exc)
        return ParsedDocument(text="", file_type="pdf", ocr_status="failed",
                              ocr_quality=None, page_count=None,
                              parser_notes=[f"PDF 解析失败：{exc}"])


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        logger.warning("python-docx not available: %s", exc)
        return ParsedDocument(text="", file_type="word", ocr_status="failed",
                              ocr_quality=None, page_count=None,
                              parser_notes=["python-docx 未安装"])
    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return ParsedDocument(
            text="\n\n".join(parts), file_type="word", ocr_status="none",
            ocr_quality=None, page_count=None, parser_notes=[],
        )
    except Exception as exc:
        logger.warning("DOCX parse failed: %s", exc)
        return ParsedDocument(text="", file_type="word", ocr_status="failed",
                              ocr_quality=None, page_count=None,
                              parser_notes=[f"Word 解析失败：{exc}"])


def _parse_markdown(text: str) -> ParsedDocument:
    try:
        import markdown as md  # type: ignore
        from bs4 import BeautifulSoup  # type: ignore
        html = md.markdown(text)
        soup = BeautifulSoup(html, "html.parser")
        plain = soup.get_text(separator="\n")
        return ParsedDocument(
            text=plain.strip(), file_type="markdown", ocr_status="none",
            ocr_quality=None, page_count=None, parser_notes=[],
        )
    except ImportError:
        # Without markdown lib, just strip basic syntax
        text = re.sub(r"^[#>\-*]+\s*", "", text, flags=re.MULTILINE)
        text = re.sub(r"`{1,3}", "", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        return ParsedDocument(
            text=text.strip(), file_type="markdown", ocr_status="none",
            ocr_quality=None, page_count=None,
            parser_notes=["markdown 库未安装，已按基础规则解析"],
        )


# --- Text utilities used by KB service ---

CJK_PATTERN = re.compile(r"[\u4e00-\u9fff]")
TOKEN_PATTERN = re.compile(r"[\u4e00-\u9fff]+|[A-Za-z][A-Za-z0-9_]*|\d+")
SENTENCE_SPLIT = re.compile(r"(?<=[。！？；\n])")


def extract_keywords(text: str, max_keywords: int = 20) -> list[str]:
    """Extract a simple keyword list from text (Chinese-aware)."""
    if not text:
        return []
    # Take CJK runs and ASCII words
    raw = TOKEN_PATTERN.findall(text)
    # Filter stopwords & short tokens
    stop = {"的", "了", "和", "是", "在", "与", "或", "及", "以及", "对", "为", "于", "由", "从", "向", "至"}
    seen: set[str] = set()
    result: list[str] = []
    for token in raw:
        token = token.strip()
        if not token or token in stop or len(token) < 2:
            continue
        if token in seen:
            continue
        seen.add(token)
        result.append(token)
        if len(result) >= max_keywords:
            break
    return result
